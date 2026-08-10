import os
import uuid
import pdfplumber
import docx
from werkzeug.utils import secure_filename
from app.utils.constants import ALLOWED_EXTENSIONS


def allowed_file(filename: str) -> bool:
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
    )


def save_upload(file, upload_folder: str) -> dict:
    """
    Saves an uploaded FileStorage object to disk.
    Returns { filename, filepath, extension }.
    """
    original_name = secure_filename(file.filename)
    ext = original_name.rsplit(".", 1)[1].lower()
    unique_name = f"{uuid.uuid4().hex}.{ext}"
    filepath = os.path.join(upload_folder, unique_name)
    file.save(filepath)
    return {
        "original_name": original_name,
        "saved_name": unique_name,
        "filepath": filepath,
        "extension": ext,
    }


def extract_text_from_file(filepath: str, extension: str) -> str:
    """
    Extract raw text from PDF, DOCX, or TXT.
    Raises ValueError for unsupported formats.
    """
    if extension == "pdf":
        return _extract_pdf(filepath)
    elif extension == "docx":
        return _extract_docx(filepath)
    elif extension == "txt":
        return _extract_txt(filepath)
    else:
        raise ValueError(f"Unsupported file extension: {extension}")


# ── private helpers ────────────────────────────────────────────────────────────

def _extract_pdf(filepath: str) -> str:
    """
    Extracts visible text AND hyperlink URIs.

    Why hyperlinks too: many resume templates show contact/profile info
    as an icon + name (e.g. an icon next to "Siddharth Srivastava" that
    links to LinkedIn) with the actual URL only present as a clickable
    link annotation — never as visible text. Plain text extraction
    alone misses this entirely, so LinkedIn/GitHub/email/phone links
    built this way were reported as "missing" even when present.
    """
    text_parts = []
    link_uris  = []

    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            for link in page.hyperlinks:
                uri = link.get("uri")
                if uri:
                    link_uris.append(uri)

    full_text = "\n".join(text_parts)
    if link_uris:
        # Appended as extra searchable text — regex extractors for
        # email/phone/linkedin/github will find matches inside these
        # URIs (e.g. "mailto:x@y.com", "https://linkedin.com/in/x")
        # exactly the same way they match visible text.
        full_text += "\n" + "\n".join(sorted(set(link_uris)))

    return full_text


def _extract_docx(filepath: str) -> str:
    """
    Extracts visible paragraph text AND hyperlink URLs.
    Same reasoning as _extract_pdf(): icon-style contact/profile links
    (icon + name, URL only in the link target) are invisible to plain
    paragraph text and were reported as "missing" even when present.
    """
    doc = docx.Document(filepath)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    link_urls = []
    for para in doc.paragraphs:
        for hyperlink in para.hyperlinks:
            if hyperlink.address:
                link_urls.append(hyperlink.address)

    full_text = "\n".join(paragraphs)
    if link_urls:
        full_text += "\n" + "\n".join(sorted(set(link_urls)))

    return full_text


def _extract_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()